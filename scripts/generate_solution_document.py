"""
Generate the Solution Architecture Word document for management review.
Output: docs/Climate_Risk_Assessment_Solution_Architecture.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent.parent
DIAGRAMS = ROOT / "docs" / "diagrams"
OUT_FILE = ROOT / "docs" / "Climate_Risk_Assessment_Solution_Architecture.docx"

NAVY = RGBColor(0x1F, 0x2D, 0x4A)
BLUE = RGBColor(0x2E, 0x5A, 0x88)
PURPLE = RGBColor(0x7A, 0x4F, 0xA3)
GREEN = RGBColor(0x3C, 0x8C, 0x6E)
AMBER = RGBColor(0xC9, 0x78, 0x2F)
GRAY = RGBColor(0x55, 0x55, 0x55)


def set_cell_background(cell, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def add_heading(doc, text, level=1, color=NAVY):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(20)
    elif level == 2:
        run.font.size = Pt(15)
    else:
        run.font.size = Pt(12.5)
    return h


def add_body(doc, text, size=10.5, bold=False, italic=False, color=None, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, bold_lead=None, size=10.5, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    if bold_lead:
        r1 = p.add_run(bold_lead)
        r1.bold = True
        r1.font.size = Pt(size)
        r2 = p.add_run(text)
        r2.font.size = Pt(size)
    else:
        r = p.add_run(text)
        r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_image(doc, filename, width_inches=6.4, caption=None):
    doc.add_picture(str(DIAGRAMS / filename), width=Inches(width_inches))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = GRAY
        cap.paragraph_format.space_after = Pt(14)


def make_table(doc, headers, rows, col_widths=None, header_color="2E5A88"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr_cells[i], header_color)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def add_page_break(doc):
    doc.add_page_break()


# =====================================================================
# BUILD DOCUMENT
# =====================================================================
doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

# ---------------------------------------------------------------
# TITLE PAGE
# ---------------------------------------------------------------
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(120)
run = title_p.add_run("Climate Risk Assessment Platform")
run.font.size = Pt(30)
run.bold = True
run.font.color.rgb = NAVY

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run("A Continuous, Hybrid AI-Powered Geospatial Risk Monitoring Data Product")
run.font.size = Pt(14)
run.font.color.rgb = BLUE
sub_p.paragraph_format.space_after = Pt(40)

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta_p.add_run("Solution Architecture & Implementation Document")
run.font.size = Pt(12.5)
run.italic = True
run.font.color.rgb = GRAY

meta_p2 = doc.add_paragraph()
meta_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta_p2.add_run("Prepared for: Engineering & Management Review")
run.font.size = Pt(10.5)
run.font.color.rgb = GRAY

add_page_break(doc)

# ---------------------------------------------------------------
# TABLE OF CONTENTS
# ---------------------------------------------------------------
add_heading(doc, "Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. Business Problem Statement",
    "3. Solution Overview",
    "4. Why This Is a Data Product",
    "5. Design Principles",
    "6. Solution Architecture",
    "7. Hybrid Architecture: Automation + Agentic AI / LLM Layer",
    "8. Data Components & Data Model",
    "9. Technology Stack",
    "10. Continuous Monitoring Data Flow",
    "11. A Day in the Life: How Users Work with the Platform",
    "12. Implementation Structure",
    "13. Business Value & Success Metrics",
    "Appendix A: Reference Design Principles (Full List)",
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

add_page_break(doc)

# ---------------------------------------------------------------
# 1. EXECUTIVE SUMMARY
# ---------------------------------------------------------------
add_heading(doc, "1. Executive Summary", level=1)
add_body(doc,
    "The Climate Risk Assessment Platform is a data product that continuously monitors wildfire "
    "and flood exposure for insured properties, replacing the industry's traditional \"assess once "
    "at underwriting\" model with real-time, evidence-based risk intelligence. The platform ingests "
    "public hazard data every 5 minutes, calculates explainable property-level risk scores, raises "
    "proactive alerts, and rolls results up into portfolio-level visibility for underwriters and risk "
    "managers.")
add_body(doc,
    "The solution uses a hybrid architecture that combines two complementary layers. A deterministic "
    "automation layer handles data ingestion, rule-based risk scoring, alerting, and monitoring — the "
    "correctness-critical, auditable core of the system. An Agentic AI / LLM layer sits on top of that "
    "core and consumes its structured output to produce contextual recommendations, portfolio "
    "insights, anomaly detection, and natural-language communication for underwriters. This division "
    "is a deliberate engineering choice: risk-critical numbers are computed by reproducible logic, "
    "while the AI layer adds judgment, synthesis, and language on top of a trustworthy foundation.")
add_body(doc,
    "This document presents the business problem, the complete solution architecture, the data "
    "product classification rationale, the hybrid automation + AI design, the underlying data model, "
    "the technology stack, and the implementation structure.")

add_page_break(doc)

# ---------------------------------------------------------------
# 2. BUSINESS PROBLEM STATEMENT
# ---------------------------------------------------------------
add_heading(doc, "2. Business Problem Statement", level=1)
add_body(doc, "The property insurance industry faces five structural limitations in how climate-related risk is assessed and managed:")

problems = [
    ("Static Risk Assessment — ", "Risk is evaluated once at policy inception and never revisited, even as environmental conditions change materially over the life of a policy."),
    ("Outdated Hazard Data — ", "Underwriting relies on static historical maps that do not reflect real-time conditions such as an active wildfire or a rapidly saturating floodplain."),
    ("Lack of Portfolio Visibility — ", "Insurers have no continuous, real-time view of exposure concentration across their book of business."),
    ("Catastrophic (CAT) Loss Exposure — ", "Poor accumulation tracking means geographic clustering of risk goes undetected until a catastrophic event forces recognition."),
    ("Limited Early Warning — ", "Customers and underwriters have no proactive mechanism to be warned of escalating risk in time to act."),
]
for lead, rest in problems:
    add_bullet(doc, rest, bold_lead=lead)

add_body(doc, "")
add_body(doc, "The Strategic Shift", bold=True, size=11.5, color=BLUE, space_after=4)
add_body(doc,
    "The platform is built around a single strategic reframing: move from assessing risk once at "
    "policy inception to continuously monitoring and updating property-level risk exposure as "
    "environmental conditions change — enabling dynamic underwriting, proactive customer engagement, "
    "and defensible, auditable pricing decisions.")

add_page_break(doc)

# ---------------------------------------------------------------
# 3. SOLUTION OVERVIEW
# ---------------------------------------------------------------
add_heading(doc, "3. Solution Overview", level=1)
add_body(doc,
    "The platform is an AI-powered geospatial risk monitoring agent that continuously assesses "
    "wildfire and flood exposure for insured properties. It is composed of six functional capabilities:")

components = [
    ("Data Ingestion Layer", "Property/exposure data plus real-time wildfire and flood hazard data from public sources (NASA FIRMS, NOAA, USGS, FEMA)."),
    ("Risk Scoring & Modeling Engine", "Dynamic, explainable risk scores combining proximity, wind-driven fire escalation, rainfall accumulation, drainage, and historical patterns."),
    ("Continuous Monitoring & Detection", "A scheduled engine (5-minute cycle, configurable) that detects environmental change and triggers threshold breaches."),
    ("Alerts & Intervention System", "Proactive notifications to underwriters, brokers, insurers, and reinsurers with recommended actions."),
    ("Portfolio Management", "Hotspot detection, geographic clustering, and scenario simulation for catastrophe planning."),
    ("Integration Layer", "Designed for connection into underwriting, claims, and pricing workflows."),
]
for name, desc in components:
    add_bullet(doc, desc, bold_lead=name + " — ")

add_page_break(doc)

# ---------------------------------------------------------------
# 4. WHY THIS IS A DATA PRODUCT
# ---------------------------------------------------------------
add_heading(doc, "4. Why This Is a Data Product", level=1)
add_body(doc,
    "This solution is deliberately classified as a data product rather than a report, a dashboard, or "
    "a one-off analysis. A data product is a system that treats data as a first-class, packaged, "
    "reusable asset — with defined inputs, defined transformations, defined consumers, and a "
    "continuous lifecycle — rather than a static output of a single analysis. This platform meets "
    "that definition on every dimension:")

dp_reasons = [
    ("Continuously Generated, Not One-Time — ", "Risk scores are recomputed on a fixed cadence (every 5 minutes) and stored as time-series snapshots, so the product is always current, not a point-in-time report."),
    ("Defined, Versioned Data Contracts — ", "Properties, hazard observations, risk assessments, and alerts each have an explicit schema (see Section 8) — the same discipline as an API contract."),
    ("Multiple Consumers, One Source of Truth — ", "The same underlying risk data serves underwriters (alerts), risk managers (portfolio dashboards), and the AI layer, without divergent copies of logic."),
    ("Reusable Compute, Not Bespoke Analysis — ", "Risk scoring, alerting, and aggregation are packaged as reusable services that apply uniformly to every property, and to properties added in the future."),
    ("Auditable and Governed — ", "Every risk score stores its contributing factors and an explanation; every alert is logged with history. This is the audit trail regulators and actuaries require."),
    ("Feeds Decisions, Not Just Dashboards — ", "Output is designed to drive action — dynamic pricing, evacuation guidance, portfolio rebalancing — which is the defining trait of a data product over a BI report."),
    ("Built for Composability — ", "The hybrid architecture explicitly treats the AI layer as a consumer of the data product's structured output, not a replacement for it — a hallmark of well-designed data products that other systems (including AI agents) can build on."),
]
for lead, rest in dp_reasons:
    add_bullet(doc, rest, bold_lead=lead)

add_body(doc, "")
add_image(doc, "04_data_product_chain.png", width_inches=6.4,
          caption="Figure 1 — Data Product Value Chain: raw hazard feeds are progressively refined into risk intelligence, AI-enriched insight, and business decisions.")

add_page_break(doc)

# ---------------------------------------------------------------
# 5. DESIGN PRINCIPLES
# ---------------------------------------------------------------
add_heading(doc, "5. Design Principles", level=1)
add_body(doc,
    "All architectural and implementation decisions are governed by twelve reference principles "
    "(full detail in Appendix A). The five most consequential to this architecture are summarized "
    "below:")

principles_summary = [
    ("Data-Driven Risk Intelligence", "Every risk score is traceable to its source data, with confidence scoring and full audit trail."),
    ("Real-Time Over Static", "Prefer continuously updated feeds over static historical hazard maps; define freshness SLAs."),
    ("Continuous Monitoring Over Point-in-Time", "Design every component assuming ongoing updates, not one-time assessment."),
    ("Transparency & Explainability", "Every score and alert must be explainable to a non-technical underwriter or a regulator."),
    ("Integration-First Architecture", "APIs and data contracts are designed from day one for underwriting/claims integration and for AI agent consumption."),
]
make_table(doc, ["Principle", "What It Drives in This Architecture"], principles_summary,
           col_widths=[2.0, 4.6])

add_page_break(doc)

# ---------------------------------------------------------------
# 6. SOLUTION ARCHITECTURE
# ---------------------------------------------------------------
add_heading(doc, "6. Solution Architecture", level=1)
add_body(doc,
    "The architecture is organized into five layers: external data sources, a scheduled ingestion "
    "layer, a persistent data layer, a processing layer (the deterministic automation core working "
    "alongside the Agentic AI layer), and stakeholder-facing outputs.")
add_image(doc, "01_high_level_architecture.png", width_inches=6.5,
          caption="Figure 2 — High-Level Solution Architecture. The deterministic automation layer (blue) computes risk-critical results; the Agentic AI / LLM layer (purple) enriches them with contextual intelligence.")

add_heading(doc, "6.1 Layer Descriptions", level=2)
layer_rows = [
    ("External Data Sources", "NASA FIRMS (wildfire), NOAA (weather), USGS (river gauges), FEMA (floodplain maps) — all public, free-tier APIs."),
    ("Data Ingestion Layer", "Scheduled fetch jobs (5-minute interval, configurable) normalize heterogeneous API responses into a common hazard schema."),
    ("Data Layer (SQLite)", "Central store for properties, hazard observations, risk assessment snapshots, alerts, and alert history — the system of record."),
    ("Deterministic Automation Layer", "Rule-based risk scoring, threshold-based alerting, continuous monitoring loop, and portfolio aggregation."),
    ("Agentic AI / LLM Layer", "Consumes structured risk output to generate recommendations, portfolio insights, anomaly detection, and natural-language alerts."),
    ("Stakeholder Outputs", "Underwriter alerts, risk manager portfolio views, and an integration surface for underwriting/claims systems."),
]
make_table(doc, ["Layer", "Responsibility"], layer_rows, col_widths=[2.0, 4.6])

add_heading(doc, "6.2 External Data Sources", level=2)
add_body(doc,
    "All hazard data is sourced from public, free-tier government and scientific APIs. Each source "
    "contributes a distinct signal to the risk model:")
source_rows = [
    ("NASA FIRMS", "Fire Information for Resource Management System — provides near-real-time satellite detections of active fires and thermal anomalies, the primary wildfire hazard signal."),
    ("NOAA", "National Oceanic and Atmospheric Administration — supplies weather data (wind speed and direction, temperature, humidity, and rainfall) used for fire-spread and flood modeling."),
    ("USGS", "United States Geological Survey — provides real-time river gauge heights and stream discharge readings used to assess flood risk near water bodies."),
    ("FEMA", "Federal Emergency Management Agency — provides official floodplain boundary maps and flood hazard zone designations used for property-level flood exposure."),
]
make_table(doc, ["Source", "Description"], source_rows, col_widths=[1.4, 5.2])

add_page_break(doc)

# ---------------------------------------------------------------
# 7. HYBRID ARCHITECTURE
# ---------------------------------------------------------------
add_heading(doc, "7. Hybrid Architecture: Automation + Agentic AI / LLM Layer", level=1)
add_body(doc,
    "A deliberate architectural decision underpins this platform: risk-critical calculations are "
    "handled by deterministic, testable automation — not by an LLM. The Agentic AI / LLM layer is "
    "designed as a hybrid add-on that consumes the automation layer's structured, validated output "
    "and adds a layer of contextual intelligence on top. This keeps the numbers underwriters rely on "
    "auditable and reproducible, while still delivering the natural-language, context-aware value that "
    "makes the platform genuinely \"AI-powered\" rather than AI-branded.")

add_image(doc, "02_hybrid_architecture.png", width_inches=6.6,
          caption="Figure 3 — Hybrid Architecture. The deterministic automation layer (top) produces structured risk data; the Agentic AI layer (bottom) consumes that data as context — it does not replace or recompute it.")

add_heading(doc, "7.1 Why Hybrid, Not AI-First", level=2)
hybrid_reasons = [
    ("Correctness & Auditability — ", "Insurance risk scores must be reproducible and defensible to regulators. Deterministic formulas with stored factor breakdowns satisfy this; LLM-generated numeric scores would not."),
    ("Cost & Latency — ", "Running an LLM call every 5 minutes across the full property set for core scoring would be costly and slow compared to arithmetic; automation is cheap to run continuously."),
    ("Reliability — ", "A monitoring system's core loop must not depend on a third-party model's availability; automation keeps the safety-critical path independent of LLM uptime."),
    ("AI Where It Adds the Most Value — ", "LLMs excel at synthesis, language, and judgment calls across multiple weak signals — exactly the layer above deterministic scoring, not the arithmetic itself."),
]
for lead, rest in hybrid_reasons:
    add_bullet(doc, rest, bold_lead=lead)

add_heading(doc, "7.2 Agentic AI Components", level=2)
agent_rows = [
    ("Recommendation Agent", "Consumes a property's risk factors and generates specific, actionable underwriter guidance (e.g., evacuation readiness, premium review)."),
    ("Portfolio Insight Agent", "Synthesizes portfolio-wide risk aggregation and hotspot data into narrative summaries and diversification recommendations."),
    ("Anomaly Detection Agent", "Flags statistically unusual hazard patterns (e.g., a 3-sigma rainfall event) that simple thresholds may under-weight."),
    ("Natural-Language Alert Generator", "Converts structured alert data into clear, audience-appropriate messages for underwriters, brokers, and customers."),
    ("Underwriting Assistant", "Combines historical loss data with current risk trajectory to support claims-prediction style reasoning."),
]
make_table(doc, ["Agent", "Function"], agent_rows, col_widths=[2.1, 4.5])

add_heading(doc, "7.3 Integration Contract", level=2)
add_body(doc,
    "Each agent receives a structured JSON context object derived directly from the risk_assessments "
    "and hazard_data tables (property details, current and prior scores, contributing factors, and "
    "recent hazard observations) and returns natural-language or structured recommendation output. "
    "This keeps the automation layer as the single source of numerical truth, with the AI layer "
    "strictly additive.")

add_page_break(doc)

# ---------------------------------------------------------------
# 8. DATA COMPONENTS & DATA MODEL
# ---------------------------------------------------------------
add_heading(doc, "8. Data Components & Data Model", level=1)
add_body(doc,
    "The system of record is a relational schema (SQLite for local deployment, portable to a managed "
    "RDBMS at scale) with six core tables.")

data_rows = [
    ("properties", "Property attributes: address, coordinates, construction type, elevation, floodplain/WUI flags."),
    ("hazard_data", "Raw ingested hazard observations from external APIs (type, source, coordinates, value, confidence, timestamps)."),
    ("risk_assessments", "Time-series snapshots of wildfire, flood, and overall risk scores with factor breakdowns and risk level classification."),
    ("alerts", "Triggered alerts with risk type, score, threshold breached, severity, and acknowledgment tracking."),
    ("alert_history", "Full audit trail of alert status changes over time."),
    ("schema_version", "Schema versioning for safe, trackable migrations."),
]
make_table(doc, ["Table", "Purpose"], data_rows, col_widths=[1.8, 4.8])

add_body(doc, "Key data quality safeguards built into the schema:", bold=True, size=10.5, space_after=4)
add_bullet(doc, "Coordinate constraints enforce valid latitude (-90 to 90) and longitude (-180 to 180) ranges at the database level.")
add_bullet(doc, "Risk scores are constrained to the 0–100 range with enumerated risk levels (low, medium, high, critical).")
add_bullet(doc, "Foreign keys link risk assessments and alerts back to properties, preserving referential integrity.")
add_bullet(doc, "Eleven indexes support fast filtering by state, coordinates, timestamp, and risk level for portfolio-scale queries.")

add_page_break(doc)

# ---------------------------------------------------------------
# 9. TECHNOLOGY STACK
# ---------------------------------------------------------------
add_heading(doc, "9. Technology Stack", level=1)
tech_rows = [
    ("Language", "Python 3.11+", "Standard for data engineering, rich geospatial ecosystem"),
    ("Database", "SQLite → portable to PostgreSQL/PostGIS at scale", "Zero-setup locally; clear scale-up path"),
    ("Geospatial", "GeoPandas, Shapely, PyProj", "Distance, proximity, and spatial calculations"),
    ("Data Processing", "Pandas, NumPy", "Time-series and property-level computation"),
    ("Scheduling", "APScheduler", "Configurable interval-based monitoring cycle"),
    ("HTTP/APIs", "Requests, Requests-Cache", "Public hazard data ingestion with caching"),
    ("Configuration", "JSON-based ConfigManager (custom)", "Dot-notation config access; no hardcoded thresholds"),
    ("Testing", "Pytest, Pytest-Cov", "Unit and integration test coverage"),
    ("Logging", "Python logging + rotating file handlers", "Full audit trail: app, alerts, and error logs"),
    ("Agentic AI Layer", "Claude API (Anthropic)", "Recommendation, insight, and natural-language generation agents"),
]
make_table(doc, ["Category", "Technology", "Rationale"], tech_rows, col_widths=[1.7, 2.4, 2.5])

add_page_break(doc)

# ---------------------------------------------------------------
# 10. CONTINUOUS MONITORING DATA FLOW
# ---------------------------------------------------------------
add_heading(doc, "10. Continuous Monitoring Data Flow", level=1)
add_body(doc,
    "The core operational loop runs on a configurable interval (5 minutes by default) and executes "
    "eight steps end-to-end, from data ingestion through portfolio aggregation, before repeating.")
add_image(doc, "03_monitoring_cycle.png", width_inches=6.6,
          caption="Figure 4 — The Continuous Monitoring Cycle. Every step is logged, and every risk change is persisted as an auditable snapshot before the cycle repeats.")

add_page_break(doc)

# ---------------------------------------------------------------
# 11. A DAY IN THE LIFE
# ---------------------------------------------------------------
add_heading(doc, "11. A Day in the Life: How Users Work with the Platform", level=1)
add_body(doc,
    "The platform is designed around its users. The three journeys below show how the same "
    "underlying data product serves different roles across a typical day — the underwriter working "
    "at the individual-policy level, the risk manager working at the portfolio level, and the broker "
    "working at the customer level. In each journey, the color of every moment indicates which part "
    "of the hybrid architecture does the work: deterministic automation for the risk-critical numbers, "
    "the Agentic AI / LLM layer for context and language, and the reporting layer for outputs and "
    "audit.")

add_heading(doc, "11.1 The Underwriter", level=2)
add_body(doc,
    "The underwriter prices and manages individual policies. The platform gives them a live, "
    "explainable risk score at the point of decision and a plain-language summary of what changed "
    "overnight — replacing static, once-at-inception assessment with continuous intelligence.")
add_image(doc, "05_day_underwriter.png", width_inches=6.7,
          caption="Figure 5 — Underwriter journey. Deterministic scores drive pricing decisions; the AI layer summarizes overnight change and recommends actions; every decision closes with an auditable record.")

add_heading(doc, "11.2 The Risk Manager", level=2)
add_body(doc,
    "The risk manager monitors accumulation and concentration risk across the whole book. The "
    "platform surfaces geographic hotspots and statistical anomalies automatically, and the AI layer "
    "turns them into narrative insight and reinsurance-planning input.")
add_image(doc, "06_day_risk_manager.png", width_inches=6.7,
          caption="Figure 6 — Risk Manager journey. Portfolio aggregation and scenario simulation run in automation; the Portfolio Insight and Anomaly Detection agents translate the numbers into decisions.")

add_heading(doc, "11.3 The Broker / Customer Liaison", level=2)
add_body(doc,
    "The broker is the customer-facing role. The platform's natural-language layer is especially "
    "valuable here: it converts technical risk changes into clear, proactive conversations with "
    "insureds about mitigation, evacuation readiness, and pricing.")
add_image(doc, "07_day_broker.png", width_inches=6.7,
          caption="Figure 7 — Broker / Customer Liaison journey. The Agentic AI layer powers proactive, plain-language customer engagement grounded in the platform's underlying risk data.")

add_page_break(doc)

# ---------------------------------------------------------------
# 12. IMPLEMENTATION STRUCTURE
# ---------------------------------------------------------------
add_heading(doc, "12. Implementation Structure", level=1)
add_body(doc,
    "The platform is implemented as a set of modular, independently testable components organized "
    "into seven functional module groups. Each module group has a clear responsibility and a defined "
    "interface to the others, which keeps the system maintainable and allows components to evolve "
    "independently.")

module_rows = [
    ("Foundation & Platform Services", "Project structure, relational database schema, configuration management, logging framework, and geospatial/time/validation utilities."),
    ("Data Ingestion", "Property/exposure data plus wildfire, weather, and flood ingestion from public APIs, with a normalization pipeline into a common hazard schema."),
    ("Risk Scoring Engine", "Wildfire and flood scoring algorithms, score aggregation, and risk-level classification — each score stored with its contributing factor breakdown."),
    ("Alerts & Continuous Monitoring", "Threshold evaluation engine, notification system, change detection, the end-to-end monitoring loop, and the interval scheduler."),
    ("Portfolio Analytics", "Portfolio-level metric aggregation, geographic hotspot detection, and reporting."),
    ("Agentic AI / LLM Layer", "Recommendation, portfolio insight, anomaly detection, and natural-language alert agents built on top of structured risk output."),
    ("Quality Assurance", "Unit and integration test suites, performance testing, and end-user/operations documentation."),
]
make_table(doc, ["Module Group", "Scope"], module_rows, col_widths=[2.2, 4.4])

add_body(doc, "")
add_body(doc,
    "Every component is configuration-driven: monitoring intervals, risk-scoring weights, and alert "
    "thresholds are defined in a central configuration file and can be tuned without code changes, "
    "supporting the Transparency and Regulatory-Readiness principles.", italic=True, color=GRAY)

add_page_break(doc)

# ---------------------------------------------------------------
# 13. BUSINESS VALUE & SUCCESS METRICS
# ---------------------------------------------------------------
add_heading(doc, "13. Business Value & Success Metrics", level=1)
add_body(doc, "The platform is designed to deliver measurable value across underwriting, portfolio management, and customer experience:")

value_rows = [
    ("Dynamic Underwriting", "Ability to reprice or flag policies mid-term based on real, current risk rather than waiting for renewal."),
    ("Reduced CAT Loss Surprise", "Continuous hotspot detection surfaces geographic accumulation risk before it becomes a catastrophic loss event."),
    ("Faster, More Confident Alerts", "5-minute (configurable) detection-to-alert cycle vs. no proactive detection in the current-state process."),
    ("Auditable Decisions", "Every score and alert is explainable and traceable to source data — reducing regulatory and dispute risk."),
    ("Scalable Architecture", "Modular services, an indexed schema, and a portable data layer are designed to scale from a local deployment to a full production book without a rewrite."),
]
make_table(doc, ["Value Driver", "Description"], value_rows, col_widths=[1.8, 4.8])

add_heading(doc, "Success Criteria", level=2)
for c in [
    "Properties continuously monitored end-to-end",
    "Live hazard data ingested every 5 minutes from public APIs",
    "Risk scores calculated, explainable, and stored as an audit trail for all properties",
    "Alerts triggered automatically when configurable thresholds are breached",
    "Portfolio-level metrics and hotspots computed and available on demand",
    "All thresholds, weights, and intervals configurable without code changes",
]:
    add_bullet(doc, c)

add_page_break(doc)

# ---------------------------------------------------------------
# APPENDIX A
# ---------------------------------------------------------------
add_heading(doc, "Appendix A: Reference Design Principles (Full List)", level=1)
full_principles = [
    ("1. Data-Driven Risk Intelligence", "All risk assessments must be grounded in real, measurable data with clear lineage and confidence scoring."),
    ("2. Real-Time Over Static", "Prioritize continuously updated data sources over static historical maps and outdated datasets."),
    ("3. Continuous Monitoring Over Point-in-Time Assessment", "Shift from assess-once-at-underwriting to monitor-continuously-throughout-policy-life."),
    ("4. Geospatial Precision", "Leverage precise coordinates and spatial relationships for granular property-level risk assessment."),
    ("5. Actionable Alerts Over Noise", "Generate alerts only when there is actionable information; avoid alert fatigue through intelligent thresholding."),
    ("6. Portfolio Visibility Through Aggregation", "Enable portfolio-level risk visibility without losing property-level granularity."),
    ("7. Integration-First Architecture", "Design for integration with underwriting, claims, and pricing workflows from day one."),
    ("8. Transparency and Explainability", "Every risk score, alert, and recommendation must be explainable to non-technical stakeholders."),
    ("9. Scalability From Day One", "Design for scale; implement incrementally without over-engineering the early build."),
    ("10. Data Quality as a First-Class Concern", "Detect, log, and surface data quality issues; degrade gracefully, never fail silently."),
    ("11. User-Centric Design", "Design all interfaces and workflows around underwriters, brokers, claims adjusters, and risk managers."),
    ("12. Regulatory and Compliance Readiness", "Design with audit trails and governance in mind from the beginning."),
]
for name, desc in full_principles:
    add_body(doc, name, bold=True, size=10.5, color=BLUE, space_after=2)
    add_body(doc, desc, size=10, space_after=8)

doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run("End of Document — Climate Risk Assessment Platform Solution Architecture")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = GRAY

doc.save(str(OUT_FILE))
print(f"[OK] Document saved to {OUT_FILE}")
